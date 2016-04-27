#encoding:utf8
from docx import Document
from docx.shared import Inches
import pandas as pd
def df2docx(df,tableName='itemName',tableLabel='itemLabel',doc=None,float_digit=0):
    if doc is None:doc=Document()
    doc.add_paragraph(u'%s: %s'%(tableName,tableLabel),style='Heading 1')
    print u'%s: %s'%(tableName,tableLabel)
    indexWidth=len(df.index[0]) if isinstance(df.index,pd.MultiIndex) else 1
    dataWidth=df.shape[1]
    tableWidth=indexWidth+dataWidth
    colHeight=len(df.columns[0]) if isinstance(df.columns,pd.MultiIndex) or isinstance(df.columns[0],(tuple,list)) else 1
    dataHeight=df.shape[0]
    tableHeight=colHeight+dataHeight
    table=doc.add_table(rows=tableHeight,cols=tableWidth,style='Table Grid')
    #将df.columns写入Table
    for j in range(df.columns.size):
        col=df.columns[j]
        if not isinstance(col,list) and not isinstance(col,tuple):col=[col]
        for i in range(colHeight):
            table.rows[i].cells[j+indexWidth].text=unicode(col[i])
    #将df.index写入到Table
    for j in range(df.index.size):
        ind=df.index[j]
        if not isinstance(ind,list) and not isinstance(ind,tuple):ind=[ind]
        for i in range(indexWidth):
            table.rows[j+colHeight].cells[i].text=unicode(ind[i])
    #将data写入到Table
    for j in range(dataHeight):
        for i in range(dataWidth):
            vl=df.values[j,i]
            if type(vl) in (float,int):
                vl=('%%.%df' % float_digit) % vl
            elif type(vl) in (unicode,str):
                pass
            else:
                print type(vl)
                raise
            table.rows[j+colHeight].cells[i+indexWidth].text=vl
    #table.columns[0].width=Inches(0.1)
    table.allow_autofit=False
    col=table.columns[1]
    for cell in col.cells:
        cell.width=Inches(6)
    #doc.save(docName)
    return doc